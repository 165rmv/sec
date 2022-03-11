# -*- coding: utf-8 -*-
from docx import Document
freqlist = ["e","a","o","s","n","r","i","l","d","t","u","c","m","p","b","h","q","y","v","g","ó","í","f","j","z","á","é","ñ","x","ú","k","w","ü"]

def decriptC():
    pass

"""
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
"""


def main(): 
    #getText("ciphered_text_caeser.docx")
    """
    f = open('ciphered_text_caeser.docx', 'r')
    document = Document(f)
    f.close()

    """
    with open('ciphered_text_caeser.docx', 'rb') as f: 
        contents = f.read()
        print(contents)



    
if __name__ == "__main__":
    main()